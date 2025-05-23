"""Microsoft Word document processor."""
import re
from typing import List, Dict, Any
from pathlib import Path
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from .base_processor import BaseProcessor, ProcessedChunk

logger = logging.getLogger(__name__)

class WordProcessor(BaseProcessor):
    """Processor for Microsoft Word documents."""
    
    def can_process(self, file_path: str) -> bool:
        """Check if this processor can handle the file."""
        return Path(file_path).suffix.lower() in ['.docx', '.doc']
        
    async def process(self, file_path: str) -> List[ProcessedChunk]:
        """Process a Word document into chunks."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        try:
            # Open document
            doc = Document(file_path)
            
            # Get document metadata
            metadata = self._extract_metadata(doc, file_path)
            
            chunks = []
            current_chunk = []
            current_length = 0
            
            # Process each paragraph
            for para_num, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue
                    
                # Get paragraph properties
                properties = self._get_paragraph_properties(paragraph)
                
                # Start new chunk for headers
                if properties["is_heading"] and current_chunk:
                    chunk_text = " ".join(current_chunk)
                    if len(chunk_text) >= self.chunk_size // 2:
                        chunk_metadata = metadata.copy()
                        chunk_metadata.update({
                            "paragraph": para_num,
                            "position": len(chunks)
                        })
                        chunks.extend(
                            self._split_into_chunks(chunk_text, chunk_metadata)
                        )
                    current_chunk = []
                    current_length = 0
                    
                # Add text to current chunk
                text = paragraph.text.strip()
                current_chunk.append(text)
                current_length += len(text) + 1
                
                # Split if chunk is too large
                if current_length >= self.chunk_size:
                    chunk_text = " ".join(current_chunk)
                    chunk_metadata = metadata.copy()
                    chunk_metadata.update({
                        "paragraph": para_num,
                        "position": len(chunks)
                    })
                    chunks.extend(
                        self._split_into_chunks(chunk_text, chunk_metadata)
                    )
                    current_chunk = []
                    current_length = 0
                    
            # Handle any remaining text
            if current_chunk:
                chunk_text = " ".join(current_chunk)
                if len(chunk_text) >= self.chunk_size // 2:
                    chunk_metadata = metadata.copy()
                    chunk_metadata.update({
                        "paragraph": len(doc.paragraphs),
                        "position": len(chunks)
                    })
                    chunks.extend(
                        self._split_into_chunks(chunk_text, chunk_metadata)
                    )
                    
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {e}")
            raise ValueError(f"Failed to process {file_path}: {e}")
            
    def _extract_metadata(
        self,
        doc: Document,
        file_path: Path
    ) -> Dict[str, Any]:
        """Extract metadata from Word document."""
        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "created": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            "modified": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
            "type": "word"
        }
        
        # Extract document properties
        core_props = doc.core_properties
        if core_props:
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.created:
                metadata["doc_created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["doc_modified"] = core_props.modified.isoformat()
                
        return metadata
        
    def _get_paragraph_properties(self, paragraph: Paragraph) -> Dict[str, Any]:
        """Get properties of a paragraph."""
        properties = {
            "is_heading": False,
            "heading_level": 0,
            "font_size": 11,  # Default size
            "is_bold": False,
            "is_italic": False
        }
        
        # Check style
        if paragraph.style:
            style_name = paragraph.style.name.lower()
            if "heading" in style_name:
                properties["is_heading"] = True
                try:
                    properties["heading_level"] = int(style_name[-1])
                except:
                    pass
                    
        # Check runs for formatting
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.font:
                if first_run.font.size:
                    properties["font_size"] = first_run.font.size.pt
                properties["is_bold"] = first_run.font.bold
                properties["is_italic"] = first_run.font.italic
                
        return properties 