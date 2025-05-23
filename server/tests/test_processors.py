"""Tests for file processors."""
import os
import pytest
from pathlib import Path
from datetime import datetime

from src.processors.base_processor import BaseProcessor, ProcessedChunk
from src.processors.markdown_processor import MarkdownProcessor
from src.processors.pdf_processor import PDFProcessor
from src.processors.word_processor import WordProcessor
from src.processors.processor_factory import ProcessorFactory

@pytest.fixture
def test_files(tmp_path):
    """Create test files for testing."""
    # Create markdown file
    md_file = tmp_path / "test.md"
    md_file.write_text("""---
title: Test Document
author: Test Author
---

# Heading 1

This is a test paragraph.

## Heading 2

Another paragraph with some **bold** and *italic* text.

### Heading 3

- List item 1
- List item 2
""")
    
    # Create Word file
    docx_file = tmp_path / "test.docx"
    from docx import Document
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph.')
    doc.add_heading('Section 1', 1)
    doc.add_paragraph('Another paragraph.')
    doc.save(str(docx_file))
    
    # Create PDF file
    pdf_file = tmp_path / "test.pdf"
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Test Document")
    page.insert_text((50, 100), "This is a test paragraph.")
    doc.save(str(pdf_file))
    doc.close()
    
    return {
        'markdown': md_file,
        'word': docx_file,
        'pdf': pdf_file
    }

def test_markdown_processor(test_files):
    """Test markdown processor."""
    processor = MarkdownProcessor()
    
    # Test file type detection
    assert processor.can_process(test_files['markdown'])
    assert not processor.can_process(test_files['word'])
    assert not processor.can_process(test_files['pdf'])
    
    # Test processing
    chunks = processor.process(test_files['markdown'])
    assert len(chunks) > 0
    
    # Check metadata
    assert chunks[0].metadata['type'] == 'markdown'
    assert chunks[0].metadata['title'] == 'Test Document'
    assert chunks[0].metadata['author'] == 'Test Author'
    
    # Check content
    text = ' '.join(chunk.content for chunk in chunks)
    assert 'Heading 1' in text
    assert 'test paragraph' in text
    assert 'List item' in text

def test_word_processor(test_files):
    """Test Word processor."""
    processor = WordProcessor()
    
    # Test file type detection
    assert processor.can_process(test_files['word'])
    assert not processor.can_process(test_files['markdown'])
    assert not processor.can_process(test_files['pdf'])
    
    # Test processing
    chunks = processor.process(test_files['word'])
    assert len(chunks) > 0
    
    # Check metadata
    assert chunks[0].metadata['type'] == 'word'
    
    # Check content
    text = ' '.join(chunk.content for chunk in chunks)
    assert 'Test Document' in text
    assert 'test paragraph' in text
    assert 'Section 1' in text

def test_pdf_processor(test_files):
    """Test PDF processor."""
    processor = PDFProcessor()
    
    # Test file type detection
    assert processor.can_process(test_files['pdf'])
    assert not processor.can_process(test_files['markdown'])
    assert not processor.can_process(test_files['word'])
    
    # Test processing
    chunks = processor.process(test_files['pdf'])
    assert len(chunks) > 0
    
    # Check metadata
    assert chunks[0].metadata['type'] == 'pdf'
    
    # Check content
    text = ' '.join(chunk.content for chunk in chunks)
    assert 'Test Document' in text
    assert 'test paragraph' in text

def test_processor_factory(test_files):
    """Test processor factory."""
    factory = ProcessorFactory()
    
    # Test file type detection
    assert factory.can_process(test_files['markdown'])
    assert factory.can_process(test_files['word'])
    assert factory.can_process(test_files['pdf'])
    assert not factory.can_process('nonexistent.xyz')
    
    # Test getting processors
    assert isinstance(factory.get_processor(test_files['markdown']), MarkdownProcessor)
    assert isinstance(factory.get_processor(test_files['word']), WordProcessor)
    assert isinstance(factory.get_processor(test_files['pdf']), PDFProcessor)
    assert factory.get_processor('nonexistent.xyz') is None
    
    # Test supported extensions
    extensions = factory.get_supported_extensions()
    assert '.md' in extensions
    assert '.docx' in extensions
    assert '.pdf' in extensions 